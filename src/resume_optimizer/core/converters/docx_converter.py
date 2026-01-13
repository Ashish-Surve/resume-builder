# src/resume_optimizer/core/converters/docx_converter.py
"""
DOCX converter for resume data.
Converts Markdown to DOCX and DOCX to PDF.
"""

import re
import logging
from typing import Optional, List
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

from ..models import ResumeData, Experience, Education
from .markdown_converter import MarkdownConverter

logger = logging.getLogger(__name__)


class DocxConverter:
    """
    Converts resume data to DOCX format.
    Supports conversion from Markdown or ResumeData directly.
    """

    # ATS-friendly fonts and sizes
    FONT_NAME = "Calibri"
    FONT_SIZE_NAME = Pt(18)
    FONT_SIZE_SECTION = Pt(12)
    FONT_SIZE_SUBSECTION = Pt(11)
    FONT_SIZE_BODY = Pt(10)

    @classmethod
    def resume_to_docx(cls, resume_data: ResumeData, output_path: Optional[Path] = None) -> bytes:
        """
        Convert ResumeData to DOCX format.

        Args:
            resume_data: The resume data to convert
            output_path: Optional path to save the file

        Returns:
            DOCX file as bytes
        """
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Add content
        cls._add_header(doc, resume_data)
        cls._add_contact_info(doc, resume_data)

        if resume_data.summary:
            cls._add_section_header(doc, "PROFESSIONAL SUMMARY")
            cls._add_summary(doc, resume_data.summary)

        if resume_data.skills:
            cls._add_section_header(doc, "TECHNICAL SKILLS")
            cls._add_skills(doc, resume_data.skills)

        if resume_data.experience:
            cls._add_section_header(doc, "PROFESSIONAL EXPERIENCE")
            cls._add_experience(doc, resume_data.experience)

        if resume_data.education:
            cls._add_section_header(doc, "EDUCATION")
            cls._add_education(doc, resume_data.education)

        if resume_data.certifications:
            cls._add_section_header(doc, "CERTIFICATIONS")
            cls._add_certifications(doc, resume_data.certifications)

        if resume_data.languages:
            cls._add_section_header(doc, "LANGUAGES")
            cls._add_languages(doc, resume_data.languages)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()

        # Also save to file if path provided
        if output_path:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

        return docx_bytes

    @classmethod
    def markdown_to_docx(cls, markdown_text: str, output_path: Optional[Path] = None) -> bytes:
        """
        Convert Markdown text to DOCX format.

        Args:
            markdown_text: Markdown formatted resume
            output_path: Optional path to save the file

        Returns:
            DOCX file as bytes
        """
        # Parse markdown to ResumeData first
        resume_data = MarkdownConverter.markdown_to_resume(markdown_text)
        return cls.resume_to_docx(resume_data, output_path)

    @classmethod
    def _add_header(cls, doc: Document, resume_data: ResumeData) -> None:
        """Add name header."""
        name = resume_data.contact_info.name or "Your Name"
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(name.upper())
        run.bold = True
        run.font.size = cls.FONT_SIZE_NAME
        run.font.name = cls.FONT_NAME

    @classmethod
    def _add_contact_info(cls, doc: Document, resume_data: ResumeData) -> None:
        """Add contact information."""
        contact = resume_data.contact_info
        contact_parts = []

        if contact.phone:
            contact_parts.append(contact.phone)
        if contact.email:
            contact_parts.append(contact.email)
        if contact.address:
            contact_parts.append(contact.address)

        if contact_parts:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(" | ".join(contact_parts))
            run.font.size = cls.FONT_SIZE_BODY
            run.font.name = cls.FONT_NAME

        # Links
        links = []
        if contact.linkedin:
            links.append(f"LinkedIn: {contact.linkedin}")
        if contact.github:
            links.append(f"GitHub: {contact.github}")

        if links:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(" | ".join(links))
            run.font.size = cls.FONT_SIZE_BODY
            run.font.name = cls.FONT_NAME

    @classmethod
    def _add_section_header(cls, doc: Document, title: str) -> None:
        """Add a section header with underline."""
        para = doc.add_paragraph()
        para.space_before = Pt(12)
        para.space_after = Pt(6)
        run = para.add_run(title)
        run.bold = True
        run.font.size = cls.FONT_SIZE_SECTION
        run.font.name = cls.FONT_NAME

        # Add horizontal line after header
        para2 = doc.add_paragraph()
        para2.paragraph_format.space_after = Pt(6)
        # Use border
        pBdr = para2._p.get_or_add_pPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        bottom = OxmlElement('w:pBdr')
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), '6')
        bottom_border.set(qn('w:space'), '1')
        bottom_border.set(qn('w:color'), '000000')
        bottom.append(bottom_border)
        pBdr.append(bottom)

    @classmethod
    def _add_summary(cls, doc: Document, summary) -> None:
        """Add professional summary."""
        # Handle dict/JSON summary
        if isinstance(summary, dict):
            for key in ['optimized_summary', 'summary', 'value']:
                if key in summary:
                    summary = summary[key]
                    break
            else:
                summary = str(next(iter(summary.values()), ''))

        para = doc.add_paragraph()
        run = para.add_run(str(summary))
        run.font.size = cls.FONT_SIZE_BODY
        run.font.name = cls.FONT_NAME
        para.paragraph_format.space_after = Pt(6)

    @classmethod
    def _add_skills(cls, doc: Document, skills: List[str]) -> None:
        """Add skills section."""
        # Join skills with bullets
        skills_text = " • ".join([s.strip() for s in skills if s and s.strip()])
        para = doc.add_paragraph()
        run = para.add_run(skills_text)
        run.font.size = cls.FONT_SIZE_BODY
        run.font.name = cls.FONT_NAME
        para.paragraph_format.space_after = Pt(6)

    @classmethod
    def _add_experience(cls, doc: Document, experiences: List[Experience]) -> None:
        """Add work experience section."""
        for exp in experiences:
            # Position and Company
            if exp.position and exp.company:
                para = doc.add_paragraph()
                run = para.add_run(exp.position)
                run.bold = True
                run.font.size = cls.FONT_SIZE_SUBSECTION
                run.font.name = cls.FONT_NAME

                # Company and dates
                company_text = f"\n{exp.company}"
                if exp.duration:
                    company_text += f" | {exp.duration}"
                elif exp.start_date or exp.end_date:
                    date_range = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                    company_text += f" | {date_range}"

                run2 = para.add_run(company_text)
                run2.font.size = cls.FONT_SIZE_BODY
                run2.font.name = cls.FONT_NAME

            # Description bullets
            for desc in exp.description:
                if desc and desc.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.25)
                    run = para.add_run(f"• {desc.strip()}")
                    run.font.size = cls.FONT_SIZE_BODY
                    run.font.name = cls.FONT_NAME

            # Add spacing after experience
            doc.add_paragraph()

    @classmethod
    def _add_education(cls, doc: Document, education: List[Education]) -> None:
        """Add education section."""
        for edu in education:
            para = doc.add_paragraph()

            # Degree and Field
            degree_text = edu.degree or "Degree"
            if edu.field:
                degree_text += f" in {edu.field}"

            run = para.add_run(degree_text)
            run.bold = True
            run.font.size = cls.FONT_SIZE_SUBSECTION
            run.font.name = cls.FONT_NAME

            # Institution and date
            inst_text = f"\n{edu.institution or ''}"
            if edu.graduation_date:
                inst_text += f" | {edu.graduation_date}"
            if edu.gpa:
                inst_text += f" | GPA: {edu.gpa}"

            run2 = para.add_run(inst_text)
            run2.font.size = cls.FONT_SIZE_BODY
            run2.font.name = cls.FONT_NAME

            # Description
            for desc in edu.description:
                if desc and desc.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.25)
                    run = para.add_run(f"• {desc.strip()}")
                    run.font.size = cls.FONT_SIZE_BODY
                    run.font.name = cls.FONT_NAME

    @classmethod
    def _add_certifications(cls, doc: Document, certifications: List[str]) -> None:
        """Add certifications section."""
        for cert in certifications:
            if cert and cert.strip():
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                run = para.add_run(f"• {cert.strip()}")
                run.font.size = cls.FONT_SIZE_BODY
                run.font.name = cls.FONT_NAME

    @classmethod
    def _add_languages(cls, doc: Document, languages: List[str]) -> None:
        """Add languages section."""
        languages_text = " • ".join([lang.strip() for lang in languages if lang and lang.strip()])
        para = doc.add_paragraph()
        run = para.add_run(languages_text)
        run.font.size = cls.FONT_SIZE_BODY
        run.font.name = cls.FONT_NAME

    @classmethod
    def docx_to_pdf(cls, docx_path: Path, pdf_path: Path) -> Path:
        """
        Convert DOCX to PDF.

        Note: This uses the system's document converter.
        On macOS: Uses the 'textutil' and 'cupsfilter' or LibreOffice
        On Windows: Uses COM automation with Word
        On Linux: Uses LibreOffice headless mode

        Args:
            docx_path: Path to the DOCX file
            pdf_path: Path for the output PDF

        Returns:
            Path to the created PDF
        """
        import subprocess
        import platform
        import shutil

        system = platform.system()
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # Try LibreOffice first (cross-platform)
        libreoffice_paths = [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
            "/usr/bin/libreoffice",  # Linux
            "/usr/bin/soffice",  # Linux alternative
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
            "soffice",  # If in PATH
        ]

        libreoffice_path = None
        for path in libreoffice_paths:
            if shutil.which(path) or Path(path).exists():
                libreoffice_path = path
                break

        if libreoffice_path:
            try:
                output_dir = pdf_path.parent
                subprocess.run([
                    libreoffice_path,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(output_dir),
                    str(docx_path)
                ], check=True, capture_output=True, timeout=60)

                # LibreOffice creates PDF with same name as input
                generated_pdf = output_dir / (docx_path.stem + ".pdf")
                if generated_pdf.exists() and generated_pdf != pdf_path:
                    shutil.move(str(generated_pdf), str(pdf_path))

                if pdf_path.exists():
                    logger.info(f"PDF created using LibreOffice: {pdf_path}")
                    return pdf_path
            except (subprocess.SubprocessError, FileNotFoundError) as e:
                logger.warning(f"LibreOffice conversion failed: {e}")

        # Fallback: Use python-docx2pdf if available
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
            if pdf_path.exists():
                logger.info(f"PDF created using docx2pdf: {pdf_path}")
                return pdf_path
        except ImportError:
            logger.warning("docx2pdf not available")
        except Exception as e:
            logger.warning(f"docx2pdf conversion failed: {e}")

        # Final fallback: use ReportLab to create PDF from resume data
        logger.warning("Using fallback PDF generation via ReportLab")
        from ..pdf_generator.generator import ATSFriendlyPDFGenerator

        # Read the DOCX and convert to ResumeData
        from ..resume_parser.parser import TextExtractor
        docx_text = TextExtractor.extract_text(docx_path)

        # Parse basic resume from DOCX
        # This is a simplified fallback - for best results use LibreOffice
        resume = ResumeData(raw_text=docx_text)

        generator = ATSFriendlyPDFGenerator()
        from ..models import OptimizationResult
        generator.generate_pdf(
            resume_data=resume,
            optimization_result=OptimizationResult(),
            output_path=pdf_path,
            applicant_name="",
            company_name=""
        )

        return pdf_path

    @classmethod
    def save_to_file(cls, resume_data: ResumeData, file_path: Path) -> Path:
        """
        Save resume as DOCX file.

        Args:
            resume_data: The resume data
            file_path: Path to save the DOCX file

        Returns:
            Path to the saved file
        """
        cls.resume_to_docx(resume_data, file_path)
        return file_path
