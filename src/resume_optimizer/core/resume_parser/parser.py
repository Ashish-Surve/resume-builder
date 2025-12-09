"""
Resume parser module using spaCy and other NLP libraries with gm fallback.
Implements Strategy pattern for different parsing approaches.
"""

import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Optional, Any
import spacy
from pypdf import PdfReader
import docx
import re
from datetime import datetime

from ...utils.exceptions import ParsingError, FileProcessingError
from ..models import ResumeData, ContactInfo, FileType, Experience, Education


class BaseResumeParser(ABC):
    """Abstract base class for resume parsers."""

    @abstractmethod
    def parse(self, file_path: Path) -> ResumeData:
        """Parse resume from file and return structured data."""
        pass


class TextExtractor:
    """Handles text extraction from different file formats."""

    @staticmethod
    def extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from PDF: {e}")

    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(str(file_path))
            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
        except Exception as e:
            raise FileProcessingError(f"Failed to extract text from DOCX: {e}")

    @staticmethod
    def extract_from_txt(file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            return file_path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            try:
                return file_path.read_text(encoding='latin-1')
            except Exception as e:
                raise FileProcessingError(f"Failed to extract text from TXT: {e}")


class ContactInfoExtractor:
    """Extracts contact information from resume text."""

    def __init__(self):
        # Enhanced patterns for better matching
        self.email_pattern = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
        self.phone_patterns = [
            re.compile(r'\b(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}\b'),  # US
            re.compile(r'\b(?:\+?91[-.\s]?)?[6-9][0-9]{9}\b'),  # India
            re.compile(r'\b(?:\+?[1-9][0-9]{0,3}[-.\s]?)?[0-9]{4,14}\b')  # General international
        ]
        self.linkedin_pattern = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9-]+/?', re.IGNORECASE)
        self.github_pattern = re.compile(r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9-]+/?', re.IGNORECASE)

    def extract(self, text: str, nlp_doc) -> ContactInfo:
        """Extract contact information from text."""
        contact = ContactInfo()

        # Extract email
        email_matches = self.email_pattern.findall(text)
        if email_matches:
            contact.email = email_matches[0]  # Take first valid email

        # Extract phone with multiple patterns
        for pattern in self.phone_patterns:
            phone_match = pattern.search(text)
            if phone_match:
                phone = re.sub(r'[^\d+]', '', phone_match.group())
                if len(phone) >= 10:  # Valid phone should have at least 10 digits
                    contact.phone = phone_match.group()
                    break

        # Extract LinkedIn
        linkedin_matches = self.linkedin_pattern.findall(text)
        if linkedin_matches:
            contact.linkedin = linkedin_matches[0]

        # Extract GitHub
        github_matches = self.github_pattern.findall(text)
        if github_matches:
            contact.github = github_matches[0]

        # Extract name using NER with better logic
        person_entities = [ent for ent in nlp_doc.ents if ent.label_ == "PERSON"]
        if person_entities:
            # Take the first person entity that appears early in the document
            for ent in person_entities[:3]:  # Check first 3 person entities
                if ent.start_char < len(text) * 0.3:  # If appears in first 30% of text
                    contact.name = ent.text.strip()
                    break
            
            if not contact.name and person_entities:
                contact.name = person_entities[0].text.strip()

        return contact


class SkillsExtractor:
    """Extracts skills from resume text with enhanced skill detection."""

    def __init__(self):
        # Focused technical skills database
        self.tech_skills = {
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'php', 'ruby', 'swift',
            'kotlin', 'scala', 'go', 'rust', 'r', 'matlab', 'shell', 'bash',
            
            # Web Technologies
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask',
            'fastapi', 'spring', 'laravel', 'jquery', 'bootstrap', 'next.js',
            
            # Databases
            'mongodb', 'mysql', 'postgresql', 'sqlite', 'redis', 'elasticsearch', 'oracle',
            
            # Cloud & DevOps
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'terraform', 'jenkins', 'git',
            
            # Data Science & ML
            'machine learning', 'deep learning', 'data science', 'artificial intelligence',
            'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'jupyter', 'spark',
            
            # Testing & Methodologies
            'selenium', 'jest', 'pytest', 'agile', 'scrum', 'devops', 'ci/cd', 'rest api'
        }
        
        # Skill variations and aliases
        self.skill_aliases = {
            'js': 'javascript',
            'ts': 'typescript',
            'py': 'python',
            'ml': 'machine learning',
            'ai': 'artificial intelligence',
            'k8s': 'kubernetes'
        }

    def extract(self, text: str, nlp_doc) -> List[str]:
        """Extract skills from text with enhanced detection."""
        text_lower = text.lower()
        found_skills = set()

        # Find technical skills
        for skill in self.tech_skills:
            # Use word boundaries for better matching
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.add(skill.title())

        # Check skill aliases
        for alias, full_skill in self.skill_aliases.items():
            pattern = r'\b' + re.escape(alias.lower()) + r'\b'
            if re.search(pattern, text_lower) and full_skill in self.tech_skills:
                found_skills.add(full_skill.title())

        return sorted(list(found_skills))


class ExperienceExtractor:
    """Extracts work experience from resume text."""
    
    def __init__(self):
        self.company_patterns = [
            r'(?:at|@)\s+([A-Z][A-Za-z\s&\.\,]+(?:Inc|LLC|Corp|Company|Ltd)?)',
            r'([A-Z][A-Za-z\s&\.\,]+(?:Inc|LLC|Corp|Company|Ltd))',
        ]
        self.position_patterns = [
            r'((?:Senior|Junior|Lead|Principal)?\s*(?:Software Engineer|Developer|Analyst|Manager|Director|Consultant))',
            r'((?:Data Scientist|Machine Learning Engineer|Backend Developer|Frontend Developer|Full Stack Developer))'
        ]
        self.duration_pattern = re.compile(r'(\d{4})\s*[-–—]\s*(\d{4}|Present|Current)', re.IGNORECASE)

    def extract_from_section(self, section_text: str) -> List[Experience]:
        """Extract experience entries from experience section text."""
        experiences = []
        
        if not section_text:
            return experiences
        
        # Split by lines and process
        lines = [line.strip() for line in section_text.split('\n') if line.strip()]
        current_exp = None
        
        for line in lines:
            # Check for company/position patterns
            company_match = self._extract_company(line)
            position_match = self._extract_position(line)
            duration_match = self.duration_pattern.search(line)
            
            # If we find indicators of a new job entry
            if company_match or position_match or duration_match:
                if current_exp:
                    experiences.append(current_exp)
                
                current_exp = Experience(
                    company=company_match or "",
                    position=position_match or "",
                    duration=duration_match.group() if duration_match else "",
                    description=[]
                )
            elif current_exp and line and not self._is_section_header(line):
                # Add as description bullet point
                current_exp.description.append(line)
        
        # Don't forget the last experience
        if current_exp:
            experiences.append(current_exp)
        
        return experiences
    
    def _extract_company(self, text: str) -> Optional[str]:
        """Extract company name from text."""
        for pattern in self.company_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1).strip()
        return None
    
    def _extract_position(self, text: str) -> Optional[str]:
        """Extract position title from text."""
        for pattern in self.position_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def _is_section_header(self, line: str) -> bool:
        """Check if line looks like a section header."""
        return len(line) < 50 and any(keyword in line.lower() for keyword in ['experience', 'education', 'skills'])


class EducationExtractor:
    """Extracts education from resume text."""
    
    def __init__(self):
        self.degree_patterns = [
            r'(Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|M\.?A\.?|B\.?A\.?)',
            r'(B\.?Tech|M\.?Tech|MBA|BBA)'
        ]
        self.institution_patterns = [
            r'(?:University|College|Institute|School)\s+of\s+([A-Za-z\s]+)',
            r'([A-Za-z\s]+(?:University|College|Institute|School))'
        ]
        self.year_pattern = re.compile(r'(\d{4})')

    def extract_from_section(self, section_text: str) -> List[Education]:
        """Extract education entries from education section text."""
        educations = []
        
        if not section_text:
            return educations
        
        lines = [line.strip() for line in section_text.split('\n') if line.strip()]
        current_edu = None
        
        for line in lines:
            degree_match = self._extract_degree(line)
            institution_match = self._extract_institution(line)
            year_match = self.year_pattern.search(line)
            
            if degree_match or institution_match or year_match:
                if current_edu:
                    educations.append(current_edu)
                
                current_edu = Education(
                    degree=degree_match or "",
                    institution=institution_match or "",
                    year=year_match.group() if year_match else "",
                    description=[]
                )
            elif current_edu and line:
                current_edu.description.append(line)
        
        if current_edu:
            educations.append(current_edu)
        
        return educations
    
    def _extract_degree(self, text: str) -> Optional[str]:
        """Extract degree from text."""
        for pattern in self.degree_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    
    def _extract_institution(self, text: str) -> Optional[str]:
        """Extract institution name from text."""
        for pattern in self.institution_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None


class SectionExtractor:
    """Enhanced section extraction with better pattern matching."""
    
    def __init__(self):
        self.section_patterns = {
            'summary': [
                r'summary', r'profile', r'objective', r'about.*me', r'professional.*summary',
                r'career.*objective', r'overview'
            ],
            'experience': [
                r'experience', r'work.*history', r'employment.*history', r'professional.*experience',
                r'work.*experience', r'employment'
            ],
            'education': [
                r'education', r'academic.*background', r'qualifications', r'degrees'
            ],
            'skills': [
                r'skills', r'technical.*skills', r'competencies', r'technologies'
            ],
            'certifications': [
                r'certifications?', r'certificates?', r'credentials', r'licenses?'
            ],
            'projects': [
                r'projects?', r'personal.*projects?', r'key.*projects?', r'portfolio'
            ]
        }
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract sections from resume text."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        sections = {}
        current_section = None
        current_content = []
        
        for line in lines:
            line_lower = line.lower()
            
            # Check if this line is a section header
            detected_section = self._detect_section_header(line_lower)
            
            if detected_section:
                # Save previous section if exists
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                current_section = detected_section
                current_content = []
            elif current_section and line.strip():
                # Add content to current section
                current_content.append(line)
        
        # Don't forget the last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _detect_section_header(self, line: str) -> Optional[str]:
        """Detect if a line is a section header."""
        # Skip very long lines (likely not headers)
        if len(line) > 50:
            return None
            
        for section, patterns in self.section_patterns.items():
            for pattern in patterns:
                if re.search(pattern, line, re.IGNORECASE):
                    return section
        
        return None


class SpacyResumeParser(BaseResumeParser):
    """Enhanced resume parser using spaCy NLP library"""

    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            raise ParsingError("spaCy English model not found. Run: python -m spacy download en_core_web_sm")

        self.text_extractor = TextExtractor()
        self.contact_extractor = ContactInfoExtractor()
        self.skills_extractor = SkillsExtractor()
        self.section_extractor = SectionExtractor()
        self.experience_extractor = ExperienceExtractor()
        self.education_extractor = EducationExtractor()
        self.logger = logging.getLogger(__name__)

    def parse(self, file_path: Path) -> ResumeData:
        """Parse resume from file and return structured data."""
        try:
            # Determine file type
            file_type = self._get_file_type(file_path)

            # Extract text
            raw_text = self._extract_text(file_path, file_type)
            
            if not raw_text.strip():
                raise ParsingError("No text could be extracted from the file")

            # Process with spaCy
            doc = self.nlp(raw_text)

            # Extract structured information
            resume_data = ResumeData(
                raw_text=raw_text,
                file_path=file_path,
                file_type=file_type
            )

            # Extract contact information
            resume_data.contact_info = self.contact_extractor.extract(raw_text, doc)

            # Extract skills
            resume_data.skills = self.skills_extractor.extract(raw_text, doc)

            # Extract sections
            sections = self.section_extractor.extract_sections(raw_text)
            
            # Extract structured experience and education
            experience_section = sections.get('experience', '')
            education_section = sections.get('education', '')
            
            # Parse experience into Experience objects
            resume_data.experience = self.experience_extractor.extract_from_section(experience_section)
            
            # Parse education into Education objects
            resume_data.education = self.education_extractor.extract_from_section(education_section)
            
            # Safely populate other resume data with extracted sections
            if hasattr(resume_data, 'summary'):
                resume_data.summary = sections.get('summary', '')
            if hasattr(resume_data, 'certifications'):
                resume_data.certifications = sections.get('certifications', '')
            if hasattr(resume_data, 'projects'):
                resume_data.projects = sections.get('projects', '')

            # Check if parsing was successful (has meaningful data)
            if self._is_parsing_successful(resume_data):
                return resume_data
            else:
                self.logger.info("SpaCy parsing yielded insufficient data, afallback.")
                raise ParsingError("Insufficient data extracted with spaCy.")

        except Exception as e:
            self.logger.error(f"Failed to parse resume: {e}")
            raise ParsingError(f"Failed to parse resume from {file_path}: {e}")

    def _is_parsing_successful(self, resume_data: ResumeData) -> bool:
        """Check if parsing extracted meaningful information."""
        # Consider parsing successful if we have at least some key information
        has_contact = (resume_data.contact_info.name or 
                      resume_data.contact_info.email or 
                      resume_data.contact_info.phone)
        has_content = (resume_data.skills or 
                      resume_data.experience or
                      resume_data.education or
                      getattr(resume_data, 'summary', ''))
        
        return has_contact and has_content
    
    def _get_file_type(self, file_path: Path) -> FileType:
        """Determine file type from extension."""
        suffix = file_path.suffix.lower()
        if suffix == '.pdf':
            return FileType.PDF
        elif suffix == '.docx':
            return FileType.DOCX
        elif suffix == '.txt':
            return FileType.TXT
        else:
            raise ParsingError(f"Unsupported file type: {suffix}")

    def _extract_text(self, file_path: Path, file_type: FileType) -> str:
        """Extract text based on file type."""
        if file_type == FileType.PDF:
            return self.text_extractor.extract_from_pdf(file_path)
        elif file_type == FileType.DOCX:
            return self.text_extractor.extract_from_docx(file_path)
        elif file_type == FileType.TXT:
            return self.text_extractor.extract_from_txt(file_path)
        else:
            raise ParsingError(f"Unsupported file type: {file_type}")
